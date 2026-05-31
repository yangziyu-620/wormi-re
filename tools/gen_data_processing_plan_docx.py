#!/usr/bin/env python3
"""Generate the WorMI VirtualHome 数据处理方案 as a .docx document.

One-off authoring script. Content is derived from the actual data pipeline
(tools/build_virtualhome_dataset*.py), the on-disk dataset
(virtualhome-realtasks-v3-20260530), and the acceptance standard report.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path("/root/WorMI/reports/virtualhome/数据处理方案.docx")

doc = Document()

# ---- base font (CJK-friendly) ----
style = doc.styles["Normal"]
style.font.name = "DejaVu Sans"
style.font.size = Pt(10.5)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "DejaVu Sans Mono"
    r.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Pt(12)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    return t


# ============================ 封面 / 标题 ============================
title = doc.add_heading("WorMI VirtualHome 数据处理方案", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("从原始 VirtualHome ActivityPrograms 到 WorMI 训练/评测就绪数据集")
r.italic = True
r.font.size = Pt(11)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "版本 1.0 ｜ 日期 2026-05-31 ｜ 分支 tmow60-noUnity-data-pipeline\n"
    "对应数据集 virtualhome-realtasks-v3-20260530"
).font.size = Pt(9)

doc.add_paragraph()

# ============================ 1. 概述 ============================
h("1. 概述与目标", 1)
para(
    "本方案描述 WorMI（World Model Implanting，ICML 2025）复现工作中 VirtualHome 数据"
    "从原始众包数据集到训练/评测就绪 JSONL 的完整处理流程、输出规范与有效性验收标准。"
)
para("处理目标可概括为三点：", bold=True)
bullet("还原 WorMI 论文设定：4 个任务族（turnon / open / puton / placein），符号化场景图观测，seen/unseen 任务与场景双重划分。")
bullet("产出可被现有训练器（Llama-3 chat 模板）直接消费的 (instruction, observation, action, next_observation) 行级数据。")
bullet("保证数据“事前有效”：在进入训练之前，通过 eval pipeline 的专家复现验收，杜绝历史上反复出现的“训练后才发现数据不可学”的失败。")

# ============================ 2. 原始数据 ============================
h("2. 原始数据集", 1)
para(
    "原始数据为 VirtualHome 论文配套的众包活动程序数据集 "
    "programs_processed_precond_nograb_morepreconds（非 JSONL，由并行的文本程序与场景图 JSON 组成）。"
)
para("2.1 目录结构", bold=True)
code(
    "programs_processed_precond_nograb_morepreconds/\n"
    "├── executable_programs/{Scene}/{batch}/fileN_M.txt   # 动作程序（带场景实例 id）\n"
    "├── init_and_final_graphs/{Scene}/{batch}/fileN_M.json # 对应的初始/终止场景图\n"
    "├── withoutconds/   # 未加 precondition 的原始程序\n"
    "├── initstate/      # 前置条件\n"
    "└── state_list/     # 逐步状态列表"
)

para("2.2 程序脚本（executable_programs/**/*.txt）", bold=True)
para("结构固定：第 1 行标题 + 第 2 行自然语言描述 + 空行 + 动作序列；每个动作形如 [ACTION] <class_name> (scene.node_id)。")
code(
    "Drink\n"
    "I walk to the cabinet and find a glass, I switch on the faucet ...\n"
    "\n"
    "[WALK] <filing_cabinet> (1.1000)\n"
    "[FIND] <cup> (1.1001)\n"
    "[GRAB] <cup> (1.1001)\n"
    "[SWITCHON] <faucet> (1.136)\n"
    "[PUTBACK] <cup> (1.1001) <sink> (1.135)\n"
    "[DRINK] <cup> (1.1001)"
)

para("2.3 环境场景图（init_and_final_graphs/**/*.json）", bold=True)
para("与同名程序配对，给出执行前后世界状态；顶层为 init_graph 与 final_graph，各由 nodes + edges 组成（单图可达数百节点、数千条边）。")
code(
    '{\n'
    '  "init_graph": {\n'
    '    "nodes": [{"id":1,"class_name":"bathroom","category":"Rooms",\n'
    '               "properties":[],"states":["CLEAN"],...}],\n'
    '    "edges": [{"from_id":2023,"relation_type":"CLOSE","to_id":199}]\n'
    '  },\n'
    '  "final_graph": {"nodes":[...], "edges":[...]}\n'
    '}'
)
bullet("node：物体/房间，含 id、class_name、category、properties、states（OPEN/CLOSED/ON/OFF/CLEAN…）。")
bullet("edge：关系三元组 from_id --relation_type--> to_id（INSIDE / ON / CLOSE / HOLDS_RH / FACING…）。")

# ============================ 3. 处理流程 ============================
h("3. 处理流程", 1)
para("整体为单向流水线：原始程序/场景图 → 任务挖掘 → 任务与场景划分 → 专家轨迹重放 → 观测渲染 → 行级落盘 → 验收。")

para("阶段 1：真实任务挖掘", bold=True)
para(
    "从 executable_programs 的终末操作动作映射为 4 个 WorMI 任务族，按真实频率选取，"
    "替代旧版“按对象类合成枚举”导致的数据塌缩。映射规则："
)
table(
    ["原始动词", "WorMI 任务族", "元数(<obj> 数)", "语义"],
    [
        ["SWITCHON", "turnon", "1", "打开电器"],
        ["OPEN", "open", "1", "打开容器/门"],
        ["PUTBACK <src> <tgt>", "puton", "2", "把物体放到表面上"],
        ["PUTIN <src> <tgt>", "placein", "2", "把物体放进容器"],
    ],
)
bullet("注意：原始 PUTON 表示“穿戴衣物”，不映射；puton 族统一由 PUTBACK 提供。")
bullet("按论文族配额 9/7/30/32 共 78 个任务选取，turnon/open/puton 施加“每源对象类多样性上限”，placein 近全量取（语料中唯一 placein 对仅约 35 个）。")

para("阶段 2：任务与场景划分", bold=True)
bullet("场景：共 20 个场景，6 个 seen / 14 个 unseen。")
bullet("任务：78 个任务按源对象类分层划分为 16 个 seen 任务 / 62 个 unseen 任务。")
bullet("由此交叉出 4 个候选池：seen-task×seen-scene、seen×unseen、unseen×seen、unseen×unseen。")

para("阶段 3：专家轨迹重放（EvolvingGraph）", bold=True)
para(
    "对每个 (task, scene)，在该场景 init_graph 上用 EvolvingGraph 的 ScriptExecutor 做 paper-like graph planner 重放，"
    "逐步执行专家动作，产生 (状态_t, 动作_t, 状态_{t+1}) 的转移序列。"
)
bullet("跳过原因被记录到 quality_report（如 execution_failed、already_satisfied、prefilter_missing:<obj>）。")
bullet("first-action collapse gate：抑制首动作过度集中，避免“一招走天下”的退化分布。")

para("阶段 4：观测渲染与行级落盘", bold=True)
para("把每一步的场景图扁平化为按字典序排列的 (object, relation, object|room) 三元组字符串，组装成行级样本：")
code(
    '{\n'
    '  "instruction": "Put food food on table",\n'
    '  "observation": "(after_shave, inside, bathroom), ..., (character, hold, food_food), ...",\n'
    '  "action": "walk table",\n'
    '  "next_observation": "(after_shave, inside, bathroom), ..., (character, close, table), ...",\n'
    '  "_meta": {"scene": "...", "split": "...", "task_family": "puton",\n'
    '            "trajectory_id": "...", "step_index": 5, "num_steps": 7,\n'
    '            "protocol": "wormi_paper_aligned_v1"}\n'
    '}'
)
bullet("核心字段（is_valid 校验）：instruction、observation、action、next_observation。")
bullet("_meta 为可选追溯信息，不参与校验，也不进入 chat 训练模板。")
bullet("训练前再由 _load_source/as_chat 派生 behavior_cloning / dynamics / affordance 三类辅助任务并转成 Llama-3 chat 模板。")

# ============================ 4. 输出格式与划分 ============================
h("4. 输出数据集结构与划分", 1)
para("最终落盘目录（virtualhome-realtasks-v3-20260530）：", bold=True)
code(
    "virtualhome-realtasks-v3-YYYYMMDD/\n"
    "├── scene_0/ ... scene_5/         # 每个 seen 场景的 world-model 训练数据 (train.jsonl + test.jsonl)\n"
    "├── eval_col_1_seen_seen/test.jsonl       # 列1：seen 任务 × seen 场景\n"
    "├── eval_col_2_seen_unseen/test.jsonl     # 列2：seen 任务 × unseen 场景\n"
    "├── eval_col_3_unseen_unseen/test.jsonl   # 列3：unseen 任务 × unseen 场景\n"
    "├── test_*_task_*_scene.jsonl     # 汇总测试集\n"
    "├── scene_inits.json              # 各场景初始图（eval 加载）\n"
    "├── virtualhome_manifest.json     # 构建清单（配额/划分/seed）\n"
    "├── quality_report.json           # 质量统计\n"
    "└── validation-*.json             # 验收结果"
)
para("动作动词分布（quality_report，示意）：", bold=True)
table(
    ["动词", "计数"],
    [["walk", "2575"], ["grab", "571"], ["open", "541"], ["putin", "297"], ["put", "274"], ["switchon", "127"]],
)

# ============================ 5. 验收标准 ============================
h("5. 数据有效性验收标准（事前 / 强制）", 1)
para(
    "核心原则：数据有效 ⇔ EVAL pipeline（而非 build pipeline）能把专家真值动作复现到约 100%。"
    "只有全部 HARD 闸通过才允许训练（GO），任一不过即 NO-GO。",
    bold=True,
)
para("A. Eval 接口完整性（历史失败的真正死因层）", bold=True)
table(
    ["ID", "准入条件", "阈值"],
    [
        ["A1", "所有行 _meta.scene 命中 eval 使用的 scene_inits", "命中率 100%"],
        ["A2", "Expert-replay SR（真值动作走完整 eval pipeline）每列", "≥ 0.99"],
        ["A3", "Gold-script_line 天花板（已绑定 instance）每列", "≥ 0.99"],
        ["A4", "fail-only-binding（gold 过、eval 绑定路径不过）", "= 0"],
        ["A5", "train/eval 观测 renderer 字符级一致（4 族全覆盖）", "mismatch = 0"],
    ],
)
para("B. 数据完整性 / 泄漏", bold=True)
table(
    ["ID", "准入条件", "阈值"],
    [
        ["B1", "train/test trajectory_id 重叠", "= 0"],
        ["B2", "train/test 整行重叠", "= 0"],
        ["B3", "seen/unseen 任务互斥、seen/unseen 场景互斥", "重叠 = 0"],
        ["B4", "build 时 ScriptExecutor 重放失败（必要非充分）", "= 0"],
    ],
)
para("C. 可学习性 / 覆盖", bold=True)
table(
    ["ID", "准入条件", "阈值"],
    [
        ["C1", "eval 各列动作动词均在 train 出现过", "verb gap = 0"],
        ["C2", "eval 各 goal 对象类在其场景图可达", "100% reachable"],
        ["C3", "agent 状态（房间+手持物）出现在所有行", "= 100%"],
        ["C4", "标签单值：同 (instruction, observation) 无冲突 action", "冲突 = 0"],
    ],
)
para(
    "强制特性：A2 这条闸任何事后补救都过不了——只有 build 的实例绑定/执行契约与 eval 真正一致才能通过，"
    "因此本标准本身就强制“根因修复，而非事后补丁”。",
    italic=True,
)

# ============================ 6. 工具清单 ============================
h("6. 工具清单", 1)
table(
    ["脚本", "职责"],
    [
        ["tools/build_virtualhome_dataset_realtasks.py", "真实任务挖掘 + 构建（v3 主构建器）"],
        ["tools/build_virtualhome_dataset_wormi.py", "父类：场景/变体加载、replay、行 schema、划分"],
        ["tools/expert_replay_vh.py", "A 类闸：eval pipeline 专家复现验收"],
        ["tools/validate_virtualhome_dataset.py", "B/C/D 部分完整性与泄漏校验"],
        ["wormi/datasets/virtualhome.py", "VirtualHomeDataset：is_valid 校验 + as_chat 模板化"],
    ],
)

# ============================ 7. 历史教训 ============================
h("7. 已知风险与历史教训", 1)
bullet("数据塌缩：旧版“按对象类合成枚举”任务导致少数源对象主导；v3 改为从真实众包程序挖掘任务予以修复。")
bullet("验收层错位：旧 validator 只测 build 自洽（整脚本一次执行），报 100%，但真正 eval pipeline 逐步绑定仅 85–89%。务必在 eval 层验收。")
bullet("renderer 不一致：曾出现 train 用 compact 观测、eval 用全图，字符不一致导致 SR≈0；A5 闸专门拦截。")
bullet("scene_inits key 不匹配：eval 传入的 scene_inits 文件与数据 _meta.scene key 零重叠会导致逐 episode KeyError；A1 闸拦截。")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run("— 本文档由数据流水线代码与验收标准报告自动汇编生成 —")
fr.italic = True
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"WROTE {OUT}  ({OUT.stat().st_size} bytes)")
